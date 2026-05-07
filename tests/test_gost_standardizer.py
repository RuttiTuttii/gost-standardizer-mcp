from __future__ import annotations

import shutil
import sys
import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm, Pt


ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "scripts"))

import gost_standardizer  # noqa: E402
import meganorm_catalog  # noqa: E402


def _make_docx(path: Path) -> None:
    document = Document()
    section = document.sections[0]
    section.left_margin = Mm(25.4)
    section.right_margin = Mm(25.4)
    section.top_margin = Mm(25.4)
    section.bottom_margin = Mm(25.4)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title.add_run("Пояснительная записка")
    title_run.font.name = "Arial"
    title_run.font.size = Pt(10)

    heading = document.add_paragraph("Введение")
    heading.paragraph_format.first_line_indent = Mm(0)
    heading.paragraph_format.space_before = Pt(0)
    heading.paragraph_format.space_after = Pt(0)

    body = document.add_paragraph("Это тестовый абзац без нужного оформления.")
    body.paragraph_format.first_line_indent = Mm(0)
    body.paragraph_format.space_before = Pt(0)
    body.paragraph_format.space_after = Pt(0)

    document.save(path)


class GostStandardizerTests(unittest.TestCase):
    def test_profile_roundtrip_uses_temp_profiles_dir(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir_name:
            temp_dir = Path(temp_dir_name)
            with patch.object(gost_standardizer, "PROFILES_DIR", temp_dir):
                saved = gost_standardizer.save_profile(
                    name="my-college",
                    preset_name="report",
                    title="My College",
                    description="Custom preset for tests",
                    kind="organization",
                    notes=["note one"],
                )
                self.assertTrue((temp_dir / "my-college.json").exists())
                loaded = gost_standardizer.load_profile("my-college")
                self.assertEqual(loaded["key"], "my-college")
                self.assertEqual(loaded["preset"]["key"], "report")
                self.assertEqual(saved["title"], "My College")

    def test_validate_compare_and_standardize_document(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir_name:
            temp_dir = Path(temp_dir_name)
            source = temp_dir / "input.docx"
            output = temp_dir / "output.docx"
            _make_docx(source)
            before = source.read_bytes()

            validation = gost_standardizer.validate_document(str(source), preset_name="report")
            self.assertEqual(validation["kind"], "validation")
            self.assertGreaterEqual(validation["summary"]["errors"], 1)
            self.assertTrue(validation["issues"])

            comparison = gost_standardizer.compare_to_preset(str(source), preset_name="report")
            self.assertEqual(comparison["kind"], "comparison")
            self.assertGreaterEqual(comparison["summary"]["differences"], 1)
            self.assertTrue(comparison["differences"])

            standardized = gost_standardizer.standardize_document(
                str(source),
                output_path=str(output),
                preset_name="report",
                overwrite=True,
            )
            self.assertEqual(standardized["kind"], "standardization")
            self.assertTrue(output.exists())
            self.assertEqual(before, source.read_bytes())

    def test_inspect_document_and_explain_preset(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir_name:
            temp_dir = Path(temp_dir_name)
            source = temp_dir / "report.docx"
            _make_docx(source)

            inspection = gost_standardizer.inspect_document(str(source))
            self.assertEqual(inspection["preset_guess"], "report")
            self.assertIn("issues", inspection)
            explanation = gost_standardizer.explain_preset(str(source))
            self.assertEqual(explanation["guessed_preset"], "report")
            self.assertIn("scores", explanation)
            self.assertIn("signals", explanation)

    def test_legacy_doc_failure_is_clear_without_converter(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir_name:
            temp_dir = Path(temp_dir_name)
            legacy = temp_dir / "legacy.doc"
            legacy.write_bytes(b"")
            with patch.object(gost_standardizer.shutil, "which", return_value=None):
                with self.assertRaises(ValueError) as ctx:
                    gost_standardizer.inspect_document(str(legacy))
            self.assertIn("LibreOffice", str(ctx.exception))

    def test_find_current_gost_uses_local_cache(self) -> None:
        cache = meganorm_catalog.load_cache()
        self.assertGreater(len(cache.get("categories", [])), 0)

        result = meganorm_catalog.find_current_gost("ГОСТ 7.32-2017", max_pages=2, limit=5)
        self.assertEqual(result["kind"], "current-gost-search")
        self.assertIn("updated_at", result["cache"])
        self.assertIsInstance(result["documents"], list)
        if result["documents"]:
            self.assertIn("confidence", result["documents"][0])
            self.assertIn("match", result["documents"][0])

    def test_search_catalog_returns_confidence(self) -> None:
        result = meganorm_catalog.search_catalog("ГОСТ", max_pages=1, limit=3)
        self.assertEqual(result["kind"], "search")
        self.assertIn("updated_at", result["cache"])
        if result["documents"]:
            self.assertIn("confidence", result["documents"][0])


if __name__ == "__main__":
    unittest.main()
