from __future__ import annotations

from contextlib import contextmanager
from dataclasses import asdict, dataclass, field
from pathlib import Path
from typing import Any
import json
import re
import shutil
import subprocess
import tempfile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Mm, Pt


ROOT_DIR = Path(__file__).resolve().parents[1]
PROFILES_DIR = ROOT_DIR / "profiles"
DEFAULT_PROFILE_KIND = "gost"
DEFAULT_PROFILE_SOURCE = "builtin"
DEFAULT_PAGE_TOLERANCE_MM = 1.0
DEFAULT_FONT_TOLERANCE_PT = 0.5
DEFAULT_RESULT_LIMIT = 50


SECTION_KEYWORDS = {
    "report": {
        "аннотация",
        "содержание",
        "введение",
        "заключение",
        "список литературы",
        "список использованных источников",
        "приложение",
    },
    "office": {
        "приказ",
        "распоряжение",
        "положение",
        "регламент",
        "инструкция",
    },
    "technical": {
        "техническое задание",
        "пояснительная записка",
        "описание",
        "требования",
        "архитектура",
    },
}


HEADING_PATTERNS = [
    re.compile(r"^\d+(?:\.\d+)*\s+\S"),
    re.compile(
        r"^(аннотация|содержание|введение|заключение|список литературы|список использованных источников|приложения?)$",
        re.I,
    ),
    re.compile(r"^(техническое задание|пояснительная записка)$", re.I),
]

LIST_PATTERNS = [
    re.compile(r"^[•\-–—]\s+\S"),
    re.compile(r"^\(?\d+[.)]\s+\S"),
    re.compile(r"^\d+\)\s+\S"),
]


@dataclass(frozen=True)
class Preset:
    key: str
    title: str
    description: str
    page_width_mm: float
    page_height_mm: float
    margin_left_mm: float
    margin_right_mm: float
    margin_top_mm: float
    margin_bottom_mm: float
    header_mm: float
    footer_mm: float
    body_font_name: str = "Times New Roman"
    body_font_size_pt: int = 14
    heading_font_size_pt: int = 14
    title_font_size_pt: int = 16
    table_font_size_pt: int = 12
    body_line_spacing: float = 1.5
    body_first_line_indent_mm: float = 12.5
    body_space_before_pt: float = 0.0
    body_space_after_pt: float = 0.0
    heading_space_before_pt: float = 12.0
    heading_space_after_pt: float = 6.0
    title_space_before_pt: float = 0.0
    title_space_after_pt: float = 12.0
    caption_space_before_pt: float = 6.0
    caption_space_after_pt: float = 6.0


@dataclass(frozen=True)
class ValidationIssue:
    severity: str
    rule_id: str
    message: str
    confidence: float
    auto_fixable: bool
    evidence: dict[str, Any] = field(default_factory=dict)
    recommendation: str = ""


PRESETS: dict[str, Preset] = {
    "report": Preset(
        key="report",
        title="GOST report",
        description="Balanced preset for reports, coursework, explanatory notes, and formal documents.",
        page_width_mm=210,
        page_height_mm=297,
        margin_left_mm=30,
        margin_right_mm=10,
        margin_top_mm=20,
        margin_bottom_mm=20,
        header_mm=10,
        footer_mm=10,
    ),
    "office": Preset(
        key="office",
        title="GOST office",
        description="Preset for office / administrative documents with a tighter page block.",
        page_width_mm=210,
        page_height_mm=297,
        margin_left_mm=20,
        margin_right_mm=10,
        margin_top_mm=20,
        margin_bottom_mm=20,
        header_mm=10,
        footer_mm=10,
        body_line_spacing=1.0,
        body_first_line_indent_mm=0.0,
    ),
    "technical": Preset(
        key="technical",
        title="GOST technical",
        description="Preset for technical docs, specs, and engineering notes.",
        page_width_mm=210,
        page_height_mm=297,
        margin_left_mm=30,
        margin_right_mm=10,
        margin_top_mm=20,
        margin_bottom_mm=20,
        header_mm=10,
        footer_mm=10,
    ),
    "legacy-college": Preset(
        key="legacy-college",
        title="Legacy college sample",
        description="A looser baseline inspired by the archive sample documents.",
        page_width_mm=210,
        page_height_mm=297,
        margin_left_mm=28,
        margin_right_mm=6,
        margin_top_mm=18,
        margin_bottom_mm=19,
        header_mm=10,
        footer_mm=10,
    ),
}


RULE_LIBRARY: dict[str, dict[str, Any]] = {
    "document.macro_enabled": {
        "severity": "info",
        "confidence": 0.95,
        "auto_fixable": False,
        "recommendation": "Macros are preserved, but review macro-enabled documents separately.",
    },
    "document.empty_body": {
        "severity": "warning",
        "confidence": 0.98,
        "auto_fixable": False,
        "recommendation": "Add visible body text before standardizing.",
    },
    "document.title_suspect": {
        "severity": "warning",
        "confidence": 0.76,
        "auto_fixable": False,
        "recommendation": "Check whether the opening lines are a title block or just body text.",
    },
    "page.size.mismatch": {
        "severity": "error",
        "confidence": 0.99,
        "auto_fixable": True,
        "recommendation": "Set the page size to the selected profile.",
    },
    "page.margin.left": {
        "severity": "error",
        "confidence": 0.99,
        "auto_fixable": True,
        "recommendation": "Set the left margin to the selected profile.",
    },
    "page.margin.right": {
        "severity": "error",
        "confidence": 0.99,
        "auto_fixable": True,
        "recommendation": "Set the right margin to the selected profile.",
    },
    "page.margin.top": {
        "severity": "error",
        "confidence": 0.99,
        "auto_fixable": True,
        "recommendation": "Set the top margin to the selected profile.",
    },
    "page.margin.bottom": {
        "severity": "error",
        "confidence": 0.99,
        "auto_fixable": True,
        "recommendation": "Set the bottom margin to the selected profile.",
    },
    "paragraph.kind.mismatch": {
        "severity": "warning",
        "confidence": 0.8,
        "auto_fixable": True,
        "recommendation": "Apply the profile paragraph formatting to the detected structure.",
    },
    "paragraph.alignment.mismatch": {
        "severity": "warning",
        "confidence": 0.9,
        "auto_fixable": True,
        "recommendation": "Align the paragraph to the profile expectation.",
    },
    "paragraph.indent.mismatch": {
        "severity": "warning",
        "confidence": 0.9,
        "auto_fixable": True,
        "recommendation": "Adjust the paragraph indentation to the selected profile.",
    },
    "paragraph.spacing.mismatch": {
        "severity": "warning",
        "confidence": 0.88,
        "auto_fixable": True,
        "recommendation": "Adjust the paragraph spacing and line spacing to the selected profile.",
    },
    "run.font.family.mismatch": {
        "severity": "warning",
        "confidence": 0.9,
        "auto_fixable": True,
        "recommendation": "Apply the profile font family to the text runs.",
    },
    "run.font.size.mismatch": {
        "severity": "warning",
        "confidence": 0.9,
        "auto_fixable": True,
        "recommendation": "Apply the profile font size to the text runs.",
    },
    "run.font.style.mismatch": {
        "severity": "warning",
        "confidence": 0.85,
        "auto_fixable": True,
        "recommendation": "Normalize bold and italic styling to the profile expectation.",
    },
    "table.paragraph.format": {
        "severity": "warning",
        "confidence": 0.9,
        "auto_fixable": True,
        "recommendation": "Normalize table paragraph formatting to the profile expectation.",
    },
}


def _preset_to_dict(preset: Preset) -> dict[str, Any]:
    return asdict(preset)


def _profile_path(name: str) -> Path:
    safe = re.sub(r"[^A-Za-z0-9._-]+", "_", name.strip()) or "profile"
    return PROFILES_DIR / f"{safe}.json"


def _profile_payload(
    *,
    key: str,
    preset: Preset,
    title: str | None = None,
    description: str | None = None,
    kind: str = DEFAULT_PROFILE_KIND,
    source: str = DEFAULT_PROFILE_SOURCE,
    notes: list[str] | None = None,
) -> dict[str, Any]:
    return {
        "key": key,
        "title": title or preset.title,
        "description": description or preset.description,
        "kind": kind,
        "source": source,
        "preset": _preset_to_dict(preset),
        "notes": notes or [],
    }


BUILTIN_PROFILES: dict[str, dict[str, Any]] = {
    key: _profile_payload(key=key, preset=preset) for key, preset in PRESETS.items()
}


def _coerce_preset(data: dict[str, Any]) -> Preset:
    allowed = {field.name for field in Preset.__dataclass_fields__.values()}
    values = {key: value for key, value in data.items() if key in allowed}
    missing = allowed - values.keys()
    if missing:
        raise ValueError(f"Profile preset is missing fields: {', '.join(sorted(missing))}")
    return Preset(**values)


def _profile_from_payload(payload: dict[str, Any]) -> dict[str, Any]:
    preset_value = payload.get("preset")
    if isinstance(preset_value, Preset):
        preset = preset_value
    elif isinstance(preset_value, dict):
        preset = _coerce_preset(preset_value)
    else:
        raise ValueError("Profile payload must include a preset mapping")

    return {
        "key": payload.get("key") or preset.key,
        "title": payload.get("title") or preset.title,
        "description": payload.get("description") or preset.description,
        "kind": payload.get("kind") or DEFAULT_PROFILE_KIND,
        "source": payload.get("source") or DEFAULT_PROFILE_SOURCE,
        "preset": _preset_to_dict(preset),
        "notes": list(payload.get("notes") or []),
    }


def _load_profile_file(path: Path) -> dict[str, Any]:
    data = json.loads(path.read_text(encoding="utf-8"))
    payload = _profile_from_payload(data)
    payload["source"] = "file"
    payload["path"] = str(path)
    return payload


def _save_profile_file(name: str, payload: dict[str, Any]) -> Path:
    PROFILES_DIR.mkdir(parents=True, exist_ok=True)
    path = _profile_path(name)
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    return path


def list_presets() -> list[dict[str, Any]]:
    return [dict(_profile_payload(key=key, preset=preset)) for key, preset in PRESETS.items()]


def list_profiles() -> list[dict[str, Any]]:
    profiles = [dict(profile) for profile in BUILTIN_PROFILES.values()]
    if PROFILES_DIR.exists():
        for file in sorted(PROFILES_DIR.glob("*.json")):
            try:
                profiles.append(_load_profile_file(file))
            except Exception:
                continue
    return profiles


def load_profile(name: str) -> dict[str, Any]:
    candidate = Path(name).expanduser()
    if candidate.exists() and candidate.is_file():
        return _load_profile_file(candidate.resolve())

    key = name.strip().lower()
    if key in BUILTIN_PROFILES:
        payload = dict(BUILTIN_PROFILES[key])
        payload["source"] = "builtin"
        return payload

    profile_file = _profile_path(name)
    if profile_file.exists():
        return _load_profile_file(profile_file)

    available = ", ".join(sorted(BUILTIN_PROFILES))
    raise ValueError(f"Unknown profile '{name}'. Available built-ins: {available}")


def save_profile(
    name: str,
    preset_name: str | None = None,
    title: str | None = None,
    description: str | None = None,
    kind: str = "organization",
    notes: list[str] | None = None,
) -> dict[str, Any]:
    preset = resolve_preset(preset_name)
    payload = _profile_payload(
        key=name,
        preset=preset,
        title=title or name,
        description=description,
        kind=kind,
        source="file",
        notes=notes,
    )
    path = _save_profile_file(name, payload)
    payload["path"] = str(path)
    return payload


def resolve_preset(name: str | None) -> Preset:
    key = (name or "report").strip().lower()
    if key not in PRESETS:
        available = ", ".join(sorted(PRESETS))
        raise ValueError(f"Unknown preset '{name}'. Available presets: {available}")
    return PRESETS[key]


def resolve_profile(name: str | None) -> dict[str, Any]:
    if not name:
        return dict(BUILTIN_PROFILES["report"])
    return load_profile(name)


def resolve_input_path(raw_path: str, base_dir: Path | None = None) -> Path:
    candidate = Path(raw_path).expanduser()
    if not candidate.is_absolute():
        candidate = (base_dir or Path.cwd()) / candidate
    candidate = candidate.resolve()
    if not candidate.exists():
        raise FileNotFoundError(f"Input file not found: {candidate}")
    if candidate.suffix.lower() not in {".docx", ".docm", ".doc"}:
        raise ValueError("Only .docx, .docm, and .doc files are supported by the current implementation")
    return candidate


def make_output_path(input_path: Path, output_path: str | None = None) -> Path:
    if output_path:
        candidate = Path(output_path).expanduser()
        if not candidate.is_absolute():
            candidate = (input_path.parent / candidate).resolve()
        if candidate.suffix.lower() not in {".docx", ".docm"}:
            candidate = candidate.with_suffix(".docx")
        return candidate
    return input_path.with_name(f"{input_path.stem}_gost.docx")


def _convert_legacy_doc(source: Path, temp_dir: Path) -> Path:
    executable = shutil.which("soffice") or shutil.which("libreoffice")
    if not executable:
        raise ValueError(
            "Legacy .doc files require LibreOffice/soffice for conversion. "
            "Install it or convert the file to .docx manually first."
        )

    temp_dir.mkdir(parents=True, exist_ok=True)
    command = [executable, "--headless", "--convert-to", "docx", "--outdir", str(temp_dir), str(source)]
    completed = subprocess.run(command, capture_output=True, text=True, check=False)
    if completed.returncode != 0:
        raise RuntimeError(
            "Failed to convert .doc to .docx with LibreOffice: "
            f"{completed.stderr.strip() or completed.stdout.strip() or 'unknown error'}"
        )

    converted = temp_dir / f"{source.stem}.docx"
    if not converted.exists():
        raise RuntimeError("LibreOffice conversion finished, but the converted .docx file was not found")
    return converted


@contextmanager
def open_document_source(raw_path: str, base_dir: Path | None = None):
    source = resolve_input_path(raw_path, base_dir=base_dir)
    if source.suffix.lower() != ".doc":
        yield source, {"source_path": str(source), "converted_from": None, "source_kind": source.suffix.lower()}
        return

    with tempfile.TemporaryDirectory(prefix="gost-standardizer-") as temp_name:
        temp_dir = Path(temp_name)
        converted = _convert_legacy_doc(source, temp_dir)
        yield converted, {
            "source_path": str(source),
            "converted_from": str(source),
            "source_kind": ".doc",
            "converted_path": str(converted),
        }


def _set_r_fonts(run, font_name: str) -> None:
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), font_name)
    r_fonts.set(qn("w:hAnsi"), font_name)
    r_fonts.set(qn("w:cs"), font_name)
    r_fonts.set(qn("w:eastAsia"), font_name)


def _set_style_font(style, font_name: str, font_size_pt: int | None = None, bold: bool | None = None) -> None:
    font = style.font
    font.name = font_name
    if font_size_pt is not None:
        font.size = Pt(font_size_pt)
    if bold is not None:
        font.bold = bold
    if style._element.rPr is not None:
        r_fonts = style._element.rPr.rFonts
        if r_fonts is None:
            r_fonts = style._element.rPr._add_rFonts()
        r_fonts.set(qn("w:ascii"), font_name)
        r_fonts.set(qn("w:hAnsi"), font_name)
        r_fonts.set(qn("w:cs"), font_name)
        r_fonts.set(qn("w:eastAsia"), font_name)


def _normalize_text(text: str) -> str:
    return re.sub(r"\s+", " ", text.replace("\xa0", " ")).strip()


def _is_in_table(paragraph) -> bool:
    element = paragraph._element
    parent = element.getparent()
    while parent is not None:
        if parent.tag.endswith("}tc"):
            return True
        parent = parent.getparent()
    return False


def _has_numbering(paragraph) -> bool:
    p_pr = paragraph._p.pPr
    return bool(p_pr is not None and p_pr.numPr is not None)


def _classify_paragraph(paragraph, index: int, non_empty_index: int) -> str:
    text = _normalize_text(paragraph.text)
    if not text:
        return "empty"

    style_name = (paragraph.style.name if paragraph.style else "").lower()
    lowered = text.lower()

    if "title" in style_name:
        return "title"
    if "heading" in style_name:
        return "heading"
    if "caption" in style_name:
        return "caption"
    if _has_numbering(paragraph):
        return "list"
    if any(pattern.match(text) for pattern in LIST_PATTERNS):
        return "list"
    if any(pattern.match(lowered) for pattern in HEADING_PATTERNS):
        return "heading"
    if non_empty_index == 1 and len(text) <= 180 and not text.endswith((".", "!", "?")):
        return "title"
    if paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER and len(text) <= 180:
        return "title"
    if len(text) <= 72 and lowered.isupper():
        return "heading"
    if lowered.startswith(("рисунок ", "таблица ")):
        return "caption"
    if len(text) <= 96 and text[:1].isupper() and lowered.count(" ") <= 6 and not text.endswith("."):
        return "heading"
    return "body"


def _length_mm(value) -> float | None:
    if value is None:
        return None
    return float(value.mm)


def _length_pt(value) -> float | None:
    if value is None:
        return None
    return float(value.pt)


def _effective_alignment(paragraph):
    if paragraph.alignment is not None:
        return paragraph.alignment
    if paragraph.style and paragraph.style.paragraph_format.alignment is not None:
        return paragraph.style.paragraph_format.alignment
    return None


def _effective_paragraph_metrics(paragraph) -> dict[str, Any]:
    fmt = paragraph.paragraph_format
    style_fmt = paragraph.style.paragraph_format if paragraph.style else None

    def _choose(attr: str):
        value = getattr(fmt, attr)
        if value is not None:
            return value
        if style_fmt is not None:
            return getattr(style_fmt, attr)
        return None

    return {
        "alignment": _effective_alignment(paragraph),
        "first_line_indent_mm": _length_mm(_choose("first_line_indent")),
        "left_indent_mm": _length_mm(_choose("left_indent")),
        "right_indent_mm": _length_mm(_choose("right_indent")),
        "space_before_pt": _length_pt(_choose("space_before")),
        "space_after_pt": _length_pt(_choose("space_after")),
        "line_spacing": _choose("line_spacing"),
    }


def _effective_run_metrics(run, paragraph) -> dict[str, Any]:
    font = run.font
    style_font = paragraph.style.font if paragraph.style else None
    return {
        "name": font.name or (style_font.name if style_font else None),
        "size_pt": _length_pt(font.size or (style_font.size if style_font else None)),
        "bold": font.bold if font.bold is not None else (style_font.bold if style_font else None),
        "italic": font.italic if font.italic is not None else (style_font.italic if style_font else None),
    }


def _paragraph_expected_metrics(kind: str, preset: Preset) -> dict[str, Any]:
    if kind == "title":
        return {
            "alignment": WD_ALIGN_PARAGRAPH.CENTER,
            "first_line_indent_mm": 0.0,
            "left_indent_mm": 0.0,
            "right_indent_mm": 0.0,
            "space_before_pt": preset.title_space_before_pt,
            "space_after_pt": preset.title_space_after_pt,
            "line_spacing": 1.0,
            "font_size_pt": preset.title_font_size_pt,
            "bold": True,
            "italic": False,
        }
    if kind == "heading":
        return {
            "alignment": WD_ALIGN_PARAGRAPH.LEFT,
            "first_line_indent_mm": 0.0,
            "left_indent_mm": 0.0,
            "right_indent_mm": 0.0,
            "space_before_pt": preset.heading_space_before_pt,
            "space_after_pt": preset.heading_space_after_pt,
            "line_spacing": 1.0,
            "font_size_pt": preset.heading_font_size_pt,
            "bold": True,
            "italic": False,
        }
    if kind == "caption":
        return {
            "alignment": WD_ALIGN_PARAGRAPH.CENTER,
            "first_line_indent_mm": 0.0,
            "left_indent_mm": 0.0,
            "right_indent_mm": 0.0,
            "space_before_pt": preset.caption_space_before_pt,
            "space_after_pt": preset.caption_space_after_pt,
            "line_spacing": 1.0,
            "font_size_pt": preset.table_font_size_pt,
            "bold": False,
            "italic": True,
        }
    if kind == "list":
        return {
            "alignment": WD_ALIGN_PARAGRAPH.LEFT,
            "first_line_indent_mm": 0.0,
            "left_indent_mm": 8.0,
            "right_indent_mm": 0.0,
            "space_before_pt": 0.0,
            "space_after_pt": 0.0,
            "line_spacing": preset.body_line_spacing,
            "font_size_pt": preset.table_font_size_pt,
            "bold": False,
            "italic": False,
        }
    return {
        "alignment": WD_ALIGN_PARAGRAPH.JUSTIFY,
        "first_line_indent_mm": preset.body_first_line_indent_mm,
        "left_indent_mm": 0.0,
        "right_indent_mm": 0.0,
        "space_before_pt": preset.body_space_before_pt,
        "space_after_pt": preset.body_space_after_pt,
        "line_spacing": preset.body_line_spacing,
        "font_size_pt": preset.body_font_size_pt,
        "bold": False,
        "italic": False,
    }


def _issue(rule_id: str, message: str, **overrides: Any) -> dict[str, Any]:
    rule = dict(RULE_LIBRARY.get(rule_id, {}))
    payload = {
        "severity": overrides.pop("severity", rule.get("severity", "warning")),
        "rule_id": rule_id,
        "message": message,
        "confidence": float(overrides.pop("confidence", rule.get("confidence", 0.8))),
        "auto_fixable": bool(overrides.pop("auto_fixable", rule.get("auto_fixable", False))),
        "evidence": overrides.pop("evidence", {}),
        "recommendation": overrides.pop("recommendation", rule.get("recommendation", "")),
    }
    payload.update(overrides)
    return payload


def _approx_equal(actual: float | None, expected: float | None, tolerance: float) -> bool:
    if actual is None or expected is None:
        return True
    return abs(actual - expected) <= tolerance


def _collect_text(document: Document) -> str:
    chunks: list[str] = []
    for paragraph in document.paragraphs:
        text = _normalize_text(paragraph.text)
        if text:
            chunks.append(text)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    text = _normalize_text(paragraph.text)
                    if text:
                        chunks.append(text)
    return "\n".join(chunks)


def _sample_paragraphs(document: Document, sample_size: int = 8) -> list[dict[str, Any]]:
    samples: list[dict[str, Any]] = []
    for index, paragraph in enumerate(document.paragraphs):
        text = _normalize_text(paragraph.text)
        if not text:
            continue
        samples.append(
            {
                "index": index,
                "text": text,
                "style": paragraph.style.name if paragraph.style else None,
                "alignment": str(_effective_alignment(paragraph)) if _effective_alignment(paragraph) is not None else None,
                "kind": _classify_paragraph(paragraph, index, len(samples) + 1),
            }
        )
        if len(samples) >= sample_size:
            break
    return samples


def guess_preset(document: Document, source_path: Path | None = None) -> str:
    text = _collect_text(document).lower()
    scores = {key: 0 for key in PRESETS}

    if source_path is not None:
        stem = source_path.stem.lower()
        if any(token in stem for token in ("tz", "тз", "technical", "spec")):
            scores["technical"] += 2
        if any(token in stem for token in ("report", "отчет", "otchet", "пз", "poyasnit")):
            scores["report"] += 2
        if any(token in stem for token in ("order", "приказ", "reglament", "instruction")):
            scores["office"] += 2

    for key, keywords in SECTION_KEYWORDS.items():
        for keyword in keywords:
            if keyword in text:
                scores[key] += 1

    return max(scores.items(), key=lambda item: (item[1], item[0] == "report"))[0]


def explain_preset(path: str, preset_name: str | None = None, sample_size: int = 8) -> dict[str, Any]:
    with open_document_source(path) as (source, source_meta):
        document = Document(str(source))
        guessed = guess_preset(document, source)
        preset_key = (preset_name or guessed).strip().lower()
        preset = resolve_preset(preset_key)

        stem = source.stem.lower()
        signals: list[dict[str, Any]] = []
        if any(token in stem for token in ("tz", "тз", "technical", "spec")):
            signals.append({"signal": "filename", "token": stem, "preset": "technical", "weight": 2})
        if any(token in stem for token in ("report", "отчет", "otchet", "пз", "poyasnit")):
            signals.append({"signal": "filename", "token": stem, "preset": "report", "weight": 2})
        if any(token in stem for token in ("order", "приказ", "reglament", "instruction")):
            signals.append({"signal": "filename", "token": stem, "preset": "office", "weight": 2})

        lowered = _collect_text(document).lower()
        for key, keywords in SECTION_KEYWORDS.items():
            for keyword in keywords:
                if keyword in lowered:
                    signals.append({"signal": "keyword", "token": keyword, "preset": key, "weight": 1})

        return {
            "path": str(source_meta["source_path"]),
            "source_kind": source_meta["source_kind"],
            "converted_from": source_meta.get("converted_from"),
            "requested_preset": preset_name,
            "preset": asdict(preset),
            "guessed_preset": guessed,
            "scores": _preset_scores(document, source),
            "signals": signals,
        }


def _preset_scores(document: Document, source_path: Path | None = None) -> dict[str, int]:
    text = _collect_text(document).lower()
    scores = {key: 0 for key in PRESETS}

    if source_path is not None:
        stem = source_path.stem.lower()
        if any(token in stem for token in ("tz", "тз", "technical", "spec")):
            scores["technical"] += 2
        if any(token in stem for token in ("report", "отчет", "otchet", "пз", "poyasnit")):
            scores["report"] += 2
        if any(token in stem for token in ("order", "приказ", "reglament", "instruction")):
            scores["office"] += 2

    for key, keywords in SECTION_KEYWORDS.items():
        for keyword in keywords:
            if keyword in text:
                scores[key] += 1
    return scores


def _validate_page_setup(document: Document, preset: Preset) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    issues: list[dict[str, Any]] = []
    matches: list[dict[str, Any]] = []

    for section_index, section in enumerate(document.sections):
        actual = {
            "page_width_mm": round(section.page_width.mm, 2),
            "page_height_mm": round(section.page_height.mm, 2),
            "left_margin_mm": round(section.left_margin.mm, 2),
            "right_margin_mm": round(section.right_margin.mm, 2),
            "top_margin_mm": round(section.top_margin.mm, 2),
            "bottom_margin_mm": round(section.bottom_margin.mm, 2),
        }
        expected = {
            "page_width_mm": preset.page_width_mm,
            "page_height_mm": preset.page_height_mm,
            "left_margin_mm": preset.margin_left_mm,
            "right_margin_mm": preset.margin_right_mm,
            "top_margin_mm": preset.margin_top_mm,
            "bottom_margin_mm": preset.margin_bottom_mm,
        }

        for key in ("page_width_mm", "page_height_mm"):
            if not _approx_equal(actual[key], expected[key], DEFAULT_PAGE_TOLERANCE_MM):
                issues.append(
                    _issue(
                        "page.size.mismatch",
                        f"Section {section_index + 1} {key} is {actual[key]} mm, expected {expected[key]} mm",
                        evidence={"section": section_index, "actual": actual, "expected": expected},
                    )
                )
                break
        else:
            matches.append(
                {
                    "rule_id": "page.size.mismatch",
                    "section": section_index,
                    "status": "match",
                    "actual": actual,
                    "expected": expected,
                }
            )

        for key, rule_id in (
            ("left_margin_mm", "page.margin.left"),
            ("right_margin_mm", "page.margin.right"),
            ("top_margin_mm", "page.margin.top"),
            ("bottom_margin_mm", "page.margin.bottom"),
        ):
            if not _approx_equal(actual[key], expected[key], DEFAULT_PAGE_TOLERANCE_MM):
                issues.append(
                    _issue(
                        rule_id,
                        f"Section {section_index + 1} {key} is {actual[key]} mm, expected {expected[key]} mm",
                        evidence={"section": section_index, "actual": actual, "expected": expected},
                    )
                )
            else:
                matches.append(
                    {
                        "rule_id": rule_id,
                        "section": section_index,
                        "status": "match",
                        "actual": actual[key],
                        "expected": expected[key],
                    }
                )

    return issues, matches


def _validate_paragraphs(
    document: Document,
    preset: Preset,
    *,
    aggressive: bool = False,
    limit: int = DEFAULT_RESULT_LIMIT,
) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    issues: list[dict[str, Any]] = []
    matches: list[dict[str, Any]] = []
    non_empty_seen = 0

    for index, paragraph in enumerate(document.paragraphs):
        text = _normalize_text(paragraph.text)
        if not text:
            continue

        non_empty_seen += 1
        kind = _classify_paragraph(paragraph, index, non_empty_seen)
        expected = _paragraph_expected_metrics(kind, preset)
        actual = _effective_paragraph_metrics(paragraph)

        if kind in {"title", "heading", "caption", "list", "body"}:
            if kind == "body" and aggressive and len(text) <= 96 and not text.endswith("."):
                kind = "heading"
                expected = _paragraph_expected_metrics(kind, preset)
            matches.append(
                {
                    "index": index,
                    "kind": kind,
                    "text": text[:160],
                    "status": "analyzed",
                }
            )

        if kind == "body" and non_empty_seen == 1 and len(text) > 120 and not text.endswith((".", "!", "?")):
            issues.append(
                _issue(
                    "document.title_suspect",
                    "The opening text looks like body text rather than a title block.",
                    evidence={"paragraph": index, "text": text[:200], "kind": kind},
                )
            )

        if kind == "empty":
            continue

        if _is_in_table(paragraph):
            table_issue = False
            if actual["alignment"] != WD_ALIGN_PARAGRAPH.LEFT:
                table_issue = True
            if not _approx_equal(actual["first_line_indent_mm"], 0.0, DEFAULT_PAGE_TOLERANCE_MM):
                table_issue = True
            if not _approx_equal(actual["left_indent_mm"], 0.0, DEFAULT_PAGE_TOLERANCE_MM):
                table_issue = True
            if not _approx_equal(actual["right_indent_mm"], 0.0, DEFAULT_PAGE_TOLERANCE_MM):
                table_issue = True
            if table_issue:
                issues.append(
                    _issue(
                        "table.paragraph.format",
                        f"Table paragraph {index + 1} needs normalized formatting.",
                        evidence={"paragraph": index, "actual": actual, "expected": expected},
                    )
                )
                if len(issues) >= limit:
                    break
            continue

        if actual["alignment"] != expected["alignment"]:
            issues.append(
                _issue(
                    "paragraph.alignment.mismatch",
                    f"Paragraph {index + 1} alignment differs from the selected profile.",
                    evidence={"paragraph": index, "actual": actual["alignment"], "expected": expected["alignment"], "kind": kind},
                )
            )

        for key in ("first_line_indent_mm", "left_indent_mm", "right_indent_mm"):
            if not _approx_equal(actual[key], expected[key], DEFAULT_PAGE_TOLERANCE_MM):
                issues.append(
                    _issue(
                        "paragraph.indent.mismatch",
                        f"Paragraph {index + 1} {key} is {actual[key]} mm, expected {expected[key]} mm",
                        evidence={"paragraph": index, "actual": actual, "expected": expected, "kind": kind},
                    )
                )
                break

        for key in ("space_before_pt", "space_after_pt"):
            if not _approx_equal(actual[key], expected[key], DEFAULT_FONT_TOLERANCE_PT):
                issues.append(
                    _issue(
                        "paragraph.spacing.mismatch",
                        f"Paragraph {index + 1} {key} is {actual[key]} pt, expected {expected[key]} pt",
                        evidence={"paragraph": index, "actual": actual, "expected": expected, "kind": kind},
                    )
                )
                break

        line_spacing = actual["line_spacing"]
        if isinstance(line_spacing, (int, float)):
            if abs(float(line_spacing) - float(expected["line_spacing"])) > 0.05:
                issues.append(
                    _issue(
                        "paragraph.spacing.mismatch",
                        f"Paragraph {index + 1} line spacing is {line_spacing}, expected {expected['line_spacing']}",
                        evidence={"paragraph": index, "actual": actual, "expected": expected, "kind": kind},
                    )
                )

        run = next((run for run in paragraph.runs if _normalize_text(run.text)), None)
        if run is None:
            continue
        run_metrics = _effective_run_metrics(run, paragraph)

        if run_metrics["name"] and run_metrics["name"].lower() != preset.body_font_name.lower():
            issues.append(
                _issue(
                    "run.font.family.mismatch",
                    f"Paragraph {index + 1} uses font family {run_metrics['name']}, expected {preset.body_font_name}.",
                    evidence={"paragraph": index, "actual": run_metrics, "expected_font": preset.body_font_name, "kind": kind},
                )
            )
        if run_metrics["size_pt"] is not None and abs(run_metrics["size_pt"] - expected["font_size_pt"]) > DEFAULT_FONT_TOLERANCE_PT:
            issues.append(
                _issue(
                    "run.font.size.mismatch",
                    f"Paragraph {index + 1} uses font size {run_metrics['size_pt']} pt, expected {expected['font_size_pt']} pt.",
                    evidence={"paragraph": index, "actual": run_metrics, "expected_font_size": expected["font_size_pt"], "kind": kind},
                )
            )
        if run_metrics["bold"] is not None and bool(run_metrics["bold"]) != bool(expected["bold"]):
            issues.append(
                _issue(
                    "run.font.style.mismatch",
                    f"Paragraph {index + 1} bold styling differs from the selected profile.",
                    evidence={"paragraph": index, "actual": run_metrics, "expected_bold": expected["bold"], "kind": kind},
                )
            )
        if run_metrics["italic"] is not None and bool(run_metrics["italic"]) != bool(expected["italic"]):
            issues.append(
                _issue(
                    "run.font.style.mismatch",
                    f"Paragraph {index + 1} italic styling differs from the selected profile.",
                    evidence={"paragraph": index, "actual": run_metrics, "expected_italic": expected["italic"], "kind": kind},
                )
            )

        if len(issues) >= limit:
            break

    return issues[:limit], matches[:limit]


def _document_statistics(document: Document) -> dict[str, Any]:
    non_empty = [paragraph for paragraph in document.paragraphs if _normalize_text(paragraph.text)]
    return {
        "paragraphs": len(document.paragraphs),
        "tables": len(document.tables),
        "inline_shapes": len(document.inline_shapes),
        "sections": len(document.sections),
        "non_empty_paragraphs": len(non_empty),
    }


def analyze_document(
    path: str,
    *,
    profile_name: str | None = None,
    sample_size: int = 8,
    aggressive: bool = False,
    limit: int = DEFAULT_RESULT_LIMIT,
) -> dict[str, Any]:
    with open_document_source(path) as (source, source_meta):
        document = Document(str(source))
        if profile_name:
            profile = resolve_profile(profile_name)
            preset = _coerce_preset(profile["preset"])
        else:
            guessed = guess_preset(document, source)
            profile = resolve_profile(guessed)
            preset = _coerce_preset(profile["preset"])

        issues: list[dict[str, Any]] = []
        notes: list[str] = []
        if source.suffix.lower() == ".docm":
            issues.append(_issue("document.macro_enabled", "Document is macro-enabled; macros are preserved but not interpreted."))
        if not any(_normalize_text(p.text) for p in document.paragraphs):
            issues.append(_issue("document.empty_body", "No body text detected in the main document paragraphs."))

        if document.paragraphs:
            first_text = _normalize_text(document.paragraphs[0].text)
            if first_text and not any(pattern.match(first_text.lower()) for pattern in HEADING_PATTERNS):
                if not first_text.endswith((".", "!", "?")) and len(first_text) > 120:
                    notes.append("The first visible paragraph looks like body text rather than a clear title block.")

        page_issues, page_matches = _validate_page_setup(document, preset)
        paragraph_issues, paragraph_matches = _validate_paragraphs(document, preset, aggressive=aggressive, limit=limit)
        issues.extend(page_issues)
        issues.extend(paragraph_issues)

        recommendations = list(
            dict.fromkeys(
                issue["recommendation"]
                for issue in issues
                if issue.get("recommendation")
            )
        )

        severity_counts: dict[str, int] = {}
        for issue in issues:
            severity_counts[issue["severity"]] = severity_counts.get(issue["severity"], 0) + 1

        return {
            "kind": "analysis",
            "path": str(source_meta["source_path"]),
            "source_kind": source_meta["source_kind"],
            "converted_from": source_meta.get("converted_from"),
            "profile": profile,
            "preset": asdict(preset),
            "statistics": _document_statistics(document),
            "summary": {
                "issues": len(issues),
                "errors": severity_counts.get("error", 0),
                "warnings": severity_counts.get("warning", 0),
                "info": severity_counts.get("info", 0),
                "auto_fixable": sum(1 for issue in issues if issue.get("auto_fixable")),
            },
            "page_matches": page_matches,
            "paragraph_matches": paragraph_matches,
            "issues": issues[:limit],
            "recommendations": recommendations,
            "notes": notes,
            "sample_paragraphs": _sample_paragraphs(document, sample_size),
        }


def inspect_document(path: str, sample_size: int = 8) -> dict[str, Any]:
    analysis = analyze_document(path, sample_size=sample_size)
    return {
        "path": analysis["path"],
        "source_kind": analysis["source_kind"],
        "converted_from": analysis["converted_from"],
        "preset_guess": analysis["profile"]["key"],
        "profile": analysis["profile"],
        "preset": analysis["preset"],
        "statistics": analysis["statistics"],
        "deviations": [issue["message"] for issue in analysis["issues"]],
        "issues": analysis["issues"],
        "recommendations": analysis["recommendations"],
        "notes": analysis["notes"],
        "sample_paragraphs": analysis["sample_paragraphs"],
    }


def validate_document(
    path: str,
    preset_name: str | None = None,
    profile_name: str | None = None,
    aggressive: bool = False,
    sample_size: int = 8,
) -> dict[str, Any]:
    target_profile = profile_name or preset_name
    analysis = analyze_document(
        path,
        profile_name=target_profile,
        sample_size=sample_size,
        aggressive=aggressive,
    )
    return {
        "kind": "validation",
        "path": analysis["path"],
        "source_kind": analysis["source_kind"],
        "converted_from": analysis["converted_from"],
        "profile": analysis["profile"],
        "preset": analysis["preset"],
        "statistics": analysis["statistics"],
        "summary": analysis["summary"],
        "issues": analysis["issues"],
        "recommendations": analysis["recommendations"],
        "notes": analysis["notes"],
        "sample_paragraphs": analysis["sample_paragraphs"],
    }


def compare_to_preset(
    path: str,
    preset_name: str | None = None,
    profile_name: str | None = None,
    aggressive: bool = False,
    sample_size: int = 8,
) -> dict[str, Any]:
    target = profile_name or preset_name
    if target is None:
        target = "report"
    profile = resolve_profile(target)
    analysis = analyze_document(
        path,
        profile_name=target,
        sample_size=sample_size,
        aggressive=aggressive,
    )
    issues = analysis["issues"]
    matches = analysis["page_matches"] + analysis["paragraph_matches"]
    return {
        "kind": "comparison",
        "path": analysis["path"],
        "source_kind": analysis["source_kind"],
        "converted_from": analysis["converted_from"],
        "profile": profile,
        "preset": analysis["preset"],
        "summary": {
            "matches": len(matches),
            "differences": len(issues),
            "auto_fixable": sum(1 for issue in issues if issue.get("auto_fixable")),
            "errors": analysis["summary"]["errors"],
            "warnings": analysis["summary"]["warnings"],
        },
        "differences": issues,
        "matches": matches,
        "statistics": analysis["statistics"],
        "recommendations": analysis["recommendations"],
        "notes": analysis["notes"],
        "sample_paragraphs": analysis["sample_paragraphs"],
    }


def _apply_paragraph_format(paragraph, kind: str, preset: Preset, inside_table: bool = False) -> None:
    fmt = paragraph.paragraph_format
    if kind == "empty":
        fmt.space_before = Pt(0)
        fmt.space_after = Pt(0)
        return

    if inside_table:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        fmt.first_line_indent = Mm(0)
        fmt.left_indent = Mm(0)
        fmt.right_indent = Mm(0)
        fmt.space_before = Pt(0)
        fmt.space_after = Pt(0)
        fmt.line_spacing = 1.0
        return

    if kind == "title":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fmt.first_line_indent = Mm(0)
        fmt.left_indent = Mm(0)
        fmt.right_indent = Mm(0)
        fmt.space_before = Pt(preset.title_space_before_pt)
        fmt.space_after = Pt(preset.title_space_after_pt)
        fmt.line_spacing = 1.0
        return

    if kind == "heading":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        fmt.first_line_indent = Mm(0)
        fmt.left_indent = Mm(0)
        fmt.right_indent = Mm(0)
        fmt.space_before = Pt(preset.heading_space_before_pt)
        fmt.space_after = Pt(preset.heading_space_after_pt)
        fmt.line_spacing = 1.0
        return

    if kind == "caption":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fmt.first_line_indent = Mm(0)
        fmt.left_indent = Mm(0)
        fmt.right_indent = Mm(0)
        fmt.space_before = Pt(preset.caption_space_before_pt)
        fmt.space_after = Pt(preset.caption_space_after_pt)
        fmt.line_spacing = 1.0
        return

    if kind == "list":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        fmt.left_indent = Mm(8)
        fmt.first_line_indent = Mm(0)
        fmt.right_indent = Mm(0)
        fmt.space_before = Pt(0)
        fmt.space_after = Pt(0)
        fmt.line_spacing = preset.body_line_spacing
        return

    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt.first_line_indent = Mm(preset.body_first_line_indent_mm)
    fmt.left_indent = Mm(0)
    fmt.right_indent = Mm(0)
    fmt.space_before = Pt(preset.body_space_before_pt)
    fmt.space_after = Pt(preset.body_space_after_pt)
    fmt.line_spacing = preset.body_line_spacing


def _apply_run_format(run, kind: str, preset: Preset) -> None:
    if not run.text:
        return

    if kind == "title":
        run.font.name = preset.body_font_name
        run.font.size = Pt(preset.title_font_size_pt)
        run.font.bold = True
        run.font.italic = False
        _set_r_fonts(run, preset.body_font_name)
        return

    if kind == "heading":
        run.font.name = preset.body_font_name
        run.font.size = Pt(preset.heading_font_size_pt)
        run.font.bold = True
        run.font.italic = False
        _set_r_fonts(run, preset.body_font_name)
        return

    if kind == "caption":
        run.font.name = preset.body_font_name
        run.font.size = Pt(preset.table_font_size_pt)
        run.font.bold = False
        run.font.italic = True
        _set_r_fonts(run, preset.body_font_name)
        return

    size = preset.table_font_size_pt if kind == "list" else preset.body_font_size_pt
    run.font.name = preset.body_font_name
    run.font.size = Pt(size)
    run.font.bold = False
    run.font.italic = False
    _set_r_fonts(run, preset.body_font_name)


def _apply_style_defaults(document: Document, preset: Preset) -> None:
    for style_name in [
        "Normal",
        "Body Text",
        "List Paragraph",
        "Caption",
        "Title",
        "Subtitle",
        "Heading 1",
        "Heading 2",
        "Heading 3",
    ]:
        try:
            style = document.styles[style_name]
        except KeyError:
            continue

        if style_name in {"Heading 1", "Heading 2", "Heading 3"}:
            _set_style_font(style, preset.body_font_name, preset.heading_font_size_pt, True)
            continue
        if style_name == "Title":
            _set_style_font(style, preset.body_font_name, preset.title_font_size_pt, True)
            continue
        if style_name == "Caption":
            _set_style_font(style, preset.body_font_name, preset.table_font_size_pt, False)
            continue
        _set_style_font(style, preset.body_font_name, preset.body_font_size_pt, False)

        if style.paragraph_format is not None:
            fmt = style.paragraph_format
            fmt.space_before = Pt(preset.body_space_before_pt)
            fmt.space_after = Pt(preset.body_space_after_pt)
            fmt.line_spacing = preset.body_line_spacing
            if style_name == "List Paragraph":
                fmt.first_line_indent = Mm(0)
                fmt.left_indent = Mm(8)
            else:
                fmt.first_line_indent = Mm(preset.body_first_line_indent_mm)
                fmt.left_indent = Mm(0)


def _set_page_setup(document: Document, preset: Preset) -> None:
    for section in document.sections:
        section.page_width = Mm(preset.page_width_mm)
        section.page_height = Mm(preset.page_height_mm)
        section.left_margin = Mm(preset.margin_left_mm)
        section.right_margin = Mm(preset.margin_right_mm)
        section.top_margin = Mm(preset.margin_top_mm)
        section.bottom_margin = Mm(preset.margin_bottom_mm)
        section.header_distance = Mm(preset.header_mm)
        section.footer_distance = Mm(preset.footer_mm)
        section.different_first_page_header_footer = False


def _document_summary(document: Document) -> dict[str, Any]:
    return {
        "paragraphs": len(document.paragraphs),
        "tables": len(document.tables),
        "inline_shapes": len(document.inline_shapes),
        "sections": len(document.sections),
    }


def standardize_document(
    path: str,
    output_path: str | None = None,
    preset_name: str | None = None,
    profile_name: str | None = None,
    overwrite: bool = False,
    aggressive: bool = False,
    fix_page_setup: bool = True,
    fix_styles: bool = True,
    fix_paragraphs: bool = True,
    fix_tables: bool = True,
) -> dict[str, Any]:
    target = profile_name or preset_name
    if target:
        profile = resolve_profile(target)
        preset = _coerce_preset(profile["preset"])
    else:
        profile = dict(BUILTIN_PROFILES["report"])
        preset = _coerce_preset(profile["preset"])

    with open_document_source(path) as (source, source_meta):
        document = Document(str(source))
        if fix_page_setup:
            _set_page_setup(document, preset)
        if fix_styles:
            _apply_style_defaults(document, preset)

        non_empty_seen = 0
        paragraph_actions: dict[str, int] = {
            "title": 0,
            "heading": 0,
            "caption": 0,
            "list": 0,
            "body": 0,
            "empty": 0,
        }

        for index, paragraph in enumerate(document.paragraphs):
            text = _normalize_text(paragraph.text)
            if text:
                non_empty_seen += 1
            kind = _classify_paragraph(paragraph, index, non_empty_seen)
            if aggressive and kind == "body" and len(text) <= 96 and not text.endswith("."):
                kind = "heading"
            paragraph_actions[kind] = paragraph_actions.get(kind, 0) + 1
            if fix_paragraphs:
                _apply_paragraph_format(paragraph, kind, preset, inside_table=False)
                for run in paragraph.runs:
                    _apply_run_format(run, kind, preset)

        table_paragraphs = 0
        if fix_tables:
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            table_paragraphs += 1
                            _apply_paragraph_format(paragraph, "body", preset, inside_table=True)
                            for run in paragraph.runs:
                                run.font.name = preset.body_font_name
                                run.font.size = Pt(preset.table_font_size_pt)
                                run.font.bold = False
                                run.font.italic = False
                                _set_r_fonts(run, preset.body_font_name)

        if source.suffix.lower() == ".docm":
            for section in document.sections:
                section.different_first_page_header_footer = False

        output = make_output_path(source, output_path)
        if output.exists() and not overwrite:
            output = output.with_name(f"{output.stem}_v2{output.suffix}")
        output.parent.mkdir(parents=True, exist_ok=True)
        document.save(str(output))

        return {
            "kind": "standardization",
            "source_path": str(source_meta["source_path"]),
            "output_path": str(output),
            "converted_from": source_meta.get("converted_from"),
            "profile": profile,
            "preset": asdict(preset),
            "fix_flags": {
                "page_setup": fix_page_setup,
                "styles": fix_styles,
                "paragraphs": fix_paragraphs,
                "tables": fix_tables,
                "aggressive": aggressive,
            },
            "paragraph_actions": paragraph_actions,
            "table_paragraphs_touched": table_paragraphs,
            "statistics": _document_summary(document),
            "notes": [
                "The document was reformatted in a copy, preserving the original source file.",
                "Complex tables, section breaks, and custom headers should still be reviewed visually.",
            ],
        }


def render_report(result: dict[str, Any]) -> str:
    return json.dumps(result, ensure_ascii=False, indent=2)
